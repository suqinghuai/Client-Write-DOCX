import os
import sys
import pandas as pd
from docx import Document
from docx.shared import Inches
import shutil
from pathlib import Path
import configparser

def get_base_dir():
    """
    获取程序基础目录，兼容打包成exe后的情况
    """
    if getattr(sys, 'frozen', False):
        # 打包成exe后的情况
        return Path(sys.executable).parent
    else:
        # 正常Python运行的情况
        return Path(__file__).parent

def parse_config():
    """
    解析配置文件，获取列映射关系
    支持二级表头层级匹配，格式：一级表头/二级表头
    """
    base_dir = get_base_dir()
    config_file = base_dir / "config2.ini"
    config = configparser.ConfigParser()
    
    # 支持二级表头的映射关系
    column_mapping = {}
    
    try:
        with open(config_file, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if '=' in line:
                    csv_column, word_column = line.split('=', 1)
                    csv_column = csv_column.strip()
                    word_column = word_column.strip()
                    
                    if csv_column not in column_mapping:
                        column_mapping[csv_column] = []
                    column_mapping[csv_column].append(word_column)
        
        print("✓ 配置文件解析完成")
        
        return column_mapping
        
    except Exception as e:
        print(f"✗ 解析配置文件时出错: {str(e)}")
        return {}

def process_csv_to_word():
    """
    将CSV文件中的构件类型数据填入Word表格
    """
    # 定义路径
    base_dir = get_base_dir()
    csv_dir = base_dir / "CSV数据处理结果"
    template_dir = base_dir / "template"
    output_dir = base_dir / "最终输出结果"
    
    # 确保输出目录存在
    output_dir.mkdir(exist_ok=True)
    
    # 查找模板文件
    template_files = list(template_dir.glob("*.docx"))
    if not template_files:
        print("✗ 错误：在template文件夹中未找到Word模板文件")
        return
    
    template_file = template_files[0]
    print(f"📄 使用模板文件: {template_file.name}")
    
    # 查找CSV文件
    csv_files = list(csv_dir.glob("*.csv"))
    if not csv_files:
        print("✗ 错误：在CSV数据处理结果文件夹中未找到CSV文件")
        return
    
    print(f"📊 找到 {len(csv_files)} 个CSV文件需要处理")
    
    for csv_file in csv_files:
        try:
            # 读取CSV文件
            print(f"\n┌─ 处理文件: {csv_file.name}")
            
            # 读取CSV文件并过滤掉Markdown分隔符行（:---）
            with open(csv_file, 'r', encoding='utf-8') as f:
                lines = f.readlines()
            
            # 过滤掉包含Markdown分隔符的行
            filtered_lines = [line for line in lines if ':---' not in line]
            
            # 如果没有过滤掉任何行，说明没有分隔符问题
            if len(filtered_lines) == len(lines):
                df = pd.read_csv(csv_file, encoding='utf-8')
            else:
                # 使用过滤后的内容创建DataFrame
                import io
                df = pd.read_csv(io.StringIO(''.join(filtered_lines)), encoding='utf-8')
                print(f"│ ✓ 已过滤 {len(lines) - len(filtered_lines)} 行Markdown分隔符")
            
            # 解析配置文件获取列映射关系
            column_mapping = parse_config()
            if not column_mapping:
                print("│ ✗ 错误：无法解析配置文件，请检查config.ini格式")
                continue
            
            # 检查CSV文件中是否包含配置文件中定义的列
            missing_columns = []
            available_columns = {}
            
            for csv_column in column_mapping.keys():
                if csv_column in df.columns:
                    available_columns[csv_column] = df[csv_column].iloc[0:].tolist()
                    print(f"│ ✓ 找到列 '{csv_column}'，包含 {len(available_columns[csv_column])} 个数据")
                else:
                    missing_columns.append(csv_column)
            
            if not available_columns:
                print(f"│ ✗ 错误：CSV文件 {csv_file.name} 中缺少所有配置的列: {list(column_mapping.keys())}")
                continue
            
            if missing_columns:
                print(f"│ ⚠ 警告：CSV文件中缺少以下列: {missing_columns}")
            
            # 复制模板文件到输出目录
            output_filename = f"output_{csv_file.stem}.docx"
            output_file = output_dir / output_filename
            shutil.copy2(template_file, output_file)
            
            # 打开Word文档
            doc = Document(output_file)
            
            # 查找表格并填充数据
            tables_filled = 0
            total_tables = len(doc.tables)
            print(f"│ 📋 发现 {total_tables} 个表格")
            
            for table_idx, table in enumerate(doc.tables, 1):
                # 为每个表格建立列映射索引
                column_indices = {}
                
                # 精确的二级表头匹配
                print(f"│   🔍 精确匹配表格{table_idx}的表头...")
                
                # 分析表头结构：扫描前两行
                header_structure = {}
                for row_idx in range(min(2, len(table.rows))):
                    row = table.rows[row_idx]
                    for cell_idx, cell in enumerate(row.cells):
                        cell_text = cell.text.strip()
                        if cell_text:
                            if row_idx not in header_structure:
                                header_structure[row_idx] = {}
                            header_structure[row_idx][cell_idx] = cell_text
                
                # 精确匹配：检查每个配置的列映射
                for csv_column, word_columns in column_mapping.items():
                    for word_column in word_columns:
                        # 检查是否包含层级分隔符（如：一级表头/二级表头）
                        if '/' in word_column:
                            # 处理二级表头：一级表头/二级表头
                            primary_header, secondary_header = word_column.split('/', 1)
                            primary_header = primary_header.strip()
                            secondary_header = secondary_header.strip()
                            
                            # 在第一行查找一级表头，在第二行查找二级表头
                            if 0 in header_structure and 1 in header_structure:
                                for col_idx, header_text in header_structure[0].items():
                                    if primary_header in header_text:
                                        # 找到一级表头后，检查其下方的二级表头
                                        if col_idx in header_structure[1]:
                                            if secondary_header in header_structure[1][col_idx]:
                                                column_indices[csv_column] = col_idx
                                                print(f"│     ✓ 二级表头匹配: '{primary_header}/{secondary_header}' -> '{csv_column}' (第{col_idx+1}列)")
                                                break
                        else:
                            # 普通表头匹配：在第一行查找
                            if 0 in header_structure:
                                for col_idx, header_text in header_structure[0].items():
                                    if word_column in header_text:
                                        column_indices[csv_column] = col_idx
                                        print(f"│     ✓ 一级表头匹配: '{word_column}' -> '{csv_column}' (第{col_idx+1}列)")
                                        break
                
                if not column_indices:
                    print(f"│     ⚠ 表格{table_idx}未找到精确匹配的列，跳过")
                    continue
                
                # 显示匹配结果
                print(f"│     📊 精确匹配到 {len(column_indices)} 个列")
                
                # 填充数据
                start_row = 2
                max_data_rows = max(len(data) for data in available_columns.values()) if available_columns else 0
                
                print(f"│     📝 填充数据: {max_data_rows} 行")
                
                for i in range(max_data_rows):
                    target_row = start_row + i
                    if target_row >= len(table.rows):
                        print(f"│     ⚠ 表格{table_idx}行数不足，无法填充第{i + 1}个数据")
                        break
                    
                    # 为当前行填充所有列的数据
                    for csv_column, col_index in column_indices.items():
                        if csv_column in available_columns and i < len(available_columns[csv_column]):
                            data_value = available_columns[csv_column][i]
                            cell = table.cell(target_row, col_index)
                            
                            # 处理空值：将nan替换为空字符串
                            if pd.isna(data_value):
                                data_value = ""
                            else:
                                data_value = str(data_value)
                            
                            # 新策略：对于已有内容的单元格，用逗号分隔并追加新数据
                            existing_text = cell.text.strip()
                            if existing_text:
                                # 如果已有内容，用逗号分隔并追加新数据
                                if data_value and data_value not in existing_text:
                                    cell.text = existing_text + ", " + data_value
                                    tables_filled += 1
                            else:
                                # 如果单元格为空，直接填充新数据
                                cell.text = data_value
                                tables_filled += 1
                
                print(f"│     ✅ 表格{table_idx}填充完成")


            
            # 保存文档
            doc.save(output_file)
            print(f"└─ ✓ 完成: {output_filename} (填充 {tables_filled} 个数据)")
            
        except Exception as e:
            print(f"└─ ✗ 处理文件 {csv_file.name} 时出错: {str(e)}")

def wait_for_key():
    """
    等待用户按任意键继续
    """
    print("\n" + "=" * 60)
    print("按任意键退出程序...")
    try:
        # Windows系统
        if os.name == 'nt':
            import msvcrt
            msvcrt.getch()
        else:
            # Linux/Mac系统
            import termios
            import tty
            import sys
            fd = sys.stdin.fileno()
            old_settings = termios.tcgetattr(fd)
            try:
                tty.setraw(sys.stdin.fileno())
                sys.stdin.read(1)
            finally:
                termios.tcsetattr(fd, termios.TCSADRAIN, old_settings)
    except (ImportError, Exception):
        # 备用方案：使用input()
        input()

def main():
    """
    主函数
    """
    print("🚀 开始处理CSV到Word表格转换...")
    print("=" * 60)
    
    try:
        process_csv_to_word()
        print("=" * 60)
        print("✅ 处理完成！")
        print(f"📁 输出文件保存在: {get_base_dir() / '最终输出结果'}")
        
    except Exception as e:
        print(f"❌ 程序执行出错: {str(e)}")
    
    # 等待用户按任意键退出
    wait_for_key()

if __name__ == "__main__":
    main()